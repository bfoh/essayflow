from app.celery_app import celery_app
from app.schemas import JobStatus, EssayOutput, HumanizationSettings
from dotenv import load_dotenv
import redis
import json
import os
import time
import random
from datetime import datetime
from typing import Optional

# Load environment variables from .env file
load_dotenv()

# Redis client for job status updates
redis_client = redis.from_url(os.getenv("REDIS_URL", "redis://localhost:6379/0"))

# Upload directory for files
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "/tmp/essays")


def update_job_status(
    job_id: str, 
    status: JobStatus, 
    progress: int, 
    message: Optional[str] = None,
    download_url: Optional[str] = None,
    error: Optional[str] = None
):
    """Update job status in Redis."""
    job_data = redis_client.get(f"job:{job_id}")
    if job_data:
        job = json.loads(job_data)
        job["status"] = status.value
        job["progress"] = progress
        job["message"] = message
        job["updated_at"] = datetime.utcnow().isoformat()
        if download_url:
            job["download_url"] = download_url
        if error:
            job["error"] = error
        redis_client.set(f"job:{job_id}", json.dumps(job), ex=86400)  # 24 hour expiry


def api_call_with_retry(client, job_id, system_prompt, user_content, max_tokens=4000, max_retries=5):
    """Helper function to make OpenAI API calls with retry logic for rate limiting."""
    # Detect critical instructions in user content and reinforce system prompt
    if "CRITICAL - MUST FOLLOW" in user_content:
        system_prompt += " CRITICAL: You must strictly follow the user's additional instructions provided in the context. These override default behaviors."

    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content}
                ],
                response_format={"type": "json_object"},
                temperature=0.7,
                max_tokens=max_tokens
            )
            return response.choices[0].message.content
        except Exception as e:
            if "rate_limit" in str(e).lower() or "429" in str(e):
                wait_time = (2 ** attempt) + random.uniform(0, 1)
                update_job_status(job_id, JobStatus.WRITING, 50, f"Rate limited, waiting {int(wait_time)}s...")
                time.sleep(wait_time)
                if attempt == max_retries - 1:
                    raise
            else:
                raise


def claude_api_call_with_retry(client, job_id, system_prompt, user_content, max_tokens=8000, max_retries=5):
    """Helper function to make Claude API calls with retry logic for rate limiting.
    
    Claude 3.5 Sonnet is used for essay generation - it's better at long-form content.
    """
    for attempt in range(max_retries):
        try:
            response = client.messages.create(
                model="claude-3-sonnet-20240229",
                max_tokens=max_tokens,
                system=system_prompt,
                messages=[
                    {"role": "user", "content": user_content}
                ]
            )
            return response.content[0].text
        except Exception as e:
            error_str = str(e).lower()
            if "rate_limit" in error_str or "429" in str(e) or "overloaded" in error_str:
                wait_time = (2 ** attempt) + random.uniform(0, 1)
                update_job_status(job_id, JobStatus.WRITING, 50, f"Rate limited, waiting {int(wait_time)}s...")
                time.sleep(wait_time)
                if attempt == max_retries - 1:
                    raise
            else:
                raise


@celery_app.task(bind=True, name="app.tasks.process_document")
def process_document(self, job_id: str, extracted_text: str):
    """
    Task to process extracted text and reference images.
    """
    try:
        update_job_status(job_id, JobStatus.EXTRACTING, 10, "Processing extracted text...")
        
        # 1. Process Reference Images (if any)
        job_data = json.loads(redis_client.get(f"job:{job_id}"))
        ref_image_count = job_data.get("ref_image_count", 0)
        image_analysis_text = ""
        
        if ref_image_count > 0:
            update_job_status(job_id, JobStatus.EXTRACTING, 15, f"Analyzing {ref_image_count} reference images...")
            import base64
            
            # Analyze images using GPT-4o Vision
            vision_prompt = "Describe this image in detail. Focus on any data, charts, text, or key visual elements that are relevant for an academic essay."
            
            # Initialize OpenAI client
            from openai import OpenAI
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            
            for i in range(ref_image_count):
                img_key = f"job:{job_id}:ref_image:{i}"
                img_bytes = redis_client.get(img_key)
                
                if img_bytes:
                    base64_image = base64.b64encode(img_bytes).decode('utf-8')
                    
                    try:
                        response = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[
                                {
                                    "role": "user",
                                    "content": [
                                        {"type": "text", "text": vision_prompt},
                                        {
                                            "type": "image_url",
                                            "image_url": {
                                                "url": f"data:image/jpeg;base64,{base64_image}"
                                            }
                                        }
                                    ]
                                }
                            ],
                            max_tokens=500
                        )
                        image_desc = response.choices[0].message.content
                        image_analysis_text += f"\n\n[Analysis of Reference Image {i+1}]:\n{image_desc}"
                    except Exception as e:
                        print(f"Error analyzing image {i}: {e}")
            
        # Combine text and image analysis
        full_content = extracted_text + "\n\n=== REFERENCE IMAGES ANALYSIS ===" + image_analysis_text
        
        # Store extracted text
        redis_client.set(f"job:{job_id}:content", full_content, ex=86400)
        
        update_job_status(job_id, JobStatus.EXTRACTING, 20, "Text processing complete")
        
        # Chain to next task
        generate_essay.delay(job_id)
        
        return {"status": "success", "word_count": len(full_content.split())}

    except Exception as e:
        print(f"Error in process_document: {str(e)}")
        update_job_status(job_id, JobStatus.FAILED, 0, f"Processing failed: {str(e)}")
        raise


@celery_app.task(bind=True, name="app.tasks.generate_essay")
def generate_essay(self, job_id: str):
    """
    Task to generate draft essay structure and content.
    Pass 1: Academic drafting.
    """
    try:
        update_job_status(job_id, JobStatus.PLANNING, 30, "Analyzing requirements...")
        
        # Get extracted content
        extracted_text_bytes = redis_client.get(f"job:{job_id}:content")
        if not extracted_text_bytes:
            raise ValueError("No extracted content found")
            
        content = extracted_text_bytes.decode("utf-8")
        
        # Get job settings
        job_data_json = redis_client.get(f"job:{job_id}")
        job_data = json.loads(job_data_json)
        additional_prompt = job_data.get("additional_prompt", "")
        
        # Construct global context string
        # IMPORTANT: Put additional prompts FIRST so they are not lost in long context
        global_context = ""
        if additional_prompt:
            global_context += f"UPPERMOST PRIORITY - USER ADDITIONAL INSTRUCTIONS:\n{additional_prompt}\n\n"
            global_context += "INSTRUCTION: The user's additional instructions above are CRITICAL. They override any conflicting information in the assignment content below.\n\n"
        
        global_context += f"Assignment Content: {content}\n\n"
        
        # Initialize OpenAI client for essay generation (Reverted due to Anthropic 404s)
        from openai import OpenAI
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
        # Step 1: Extract requirements including word count (Claude is good at this too)
        update_job_status(job_id, JobStatus.RESEARCHING, 35, "Extracting word count requirements...")
        
        extraction_prompt = """You are an academic assignment analyzer. Extract key requirements from the assignment.

        IMPORTANT: Return ONLY a valid JSON object with no additional text. The format must be:
        {
            "required_word_count": <number>,
            "topic": "<string>",
            "key_requirements": ["<requirement1>", "<requirement2>"],
            "suggested_sections": ["<section1>", "<section2>", "<section3>", "<section4>"],
            "academic_level": "<string>",
            "citation_style": "<string>"
        }
        
        CRITICAL: Look for word counts like "2000 words", "2,000-2,500 words", "minimum 1500 words".
        If no word count is specified, default to 2000 words.
        Suggest 4-6 logical body sections based on the topic."""
        
        requirements_response = api_call_with_retry(
            client, job_id, extraction_prompt, global_context, max_tokens=1000
        )
        
        try:
            requirements = json.loads(requirements_response)
        except:
            requirements = {
                "required_word_count": 2000,
                "topic": "Essay",
                "key_requirements": [],
                "suggested_sections": ["Introduction", "Literature Review", "Analysis", "Discussion", "Conclusion"],
                "academic_level": "undergraduate",
                "citation_style": "APA"
            }
        
        # ADDITIONAL INSTRUCTIONS TAKE PRECEDENCE
        # Parse word count override from additional_prompt
        import re
        target_word_count = requirements.get("required_word_count", 2000)
        
        if additional_prompt:
            # Look for word count overrides in additional instructions
            word_count_patterns = [
                r'(\d{1,2},?\d{3})\s*words?',  # "2,000 words" or "2000 words"
                r'word\s*count[:\s]+(\d+)',     # "word count: 1500"
                r'(\d+)\s*word\s*count',        # "1500 word count"
                r'minimum\s*(\d+)',             # "minimum 1500"
                r'at\s*least\s*(\d+)',          # "at least 2000"
            ]
            for pattern in word_count_patterns:
                match = re.search(pattern, additional_prompt.lower().replace(',', ''))
                if match:
                    target_word_count = int(match.group(1).replace(',', ''))
                    break
            
            # FORCE INJECTION: Add additional instructions to key_requirements
            # This ensures downstream steps explicitly see them as requirements
            if "key_requirements" not in requirements:
                requirements["key_requirements"] = []
            requirements["key_requirements"].append(f"USER PRIORITY INSTRUCTION: {additional_prompt}")
        
        # Filter out 'Conclusion' and 'References' from body sections to prevent duplicates
        # Also handle numbered sections like "5. Conclusion"
        sections = [
            s for s in requirements.get("suggested_sections", ["Background", "Analysis", "Discussion", "Implications"]) 
            if "conclusion" not in s.lower() and "reference" not in s.lower()
        ]
        
        # Calculate words per section (intro + body sections + conclusion)
        num_sections = len(sections)
        intro_words = int(target_word_count * 0.10)  # 10% for intro
        conclusion_words = int(target_word_count * 0.10)  # 10% for conclusion
        body_words = target_word_count - intro_words - conclusion_words
        words_per_section = body_words // num_sections if num_sections > 0 else 500
        
        update_job_status(job_id, JobStatus.WRITING, 40, f"Generating {target_word_count}-word essay...")
        
        # Step 2: Generate introduction with thesis
        intro_prompt = f"""You are an expert academic writer. Write a compelling introduction for an essay.

        Topic: {requirements.get('topic', 'the given topic')}
        Target Length: {intro_words} words (CRITICAL: write at least {intro_words} words)
        Academic Level: {requirements.get('academic_level', 'undergraduate')}
        
        Requirements:
        - Include a clear, arguable thesis statement
        - Set up the key arguments that will be discussed
        - Hook the reader with an engaging opening
        - Maintain academic tone throughout
        
        Key requirements from assignment:
        {chr(10).join(requirements.get('key_requirements', []))}
        
        IMPORTANT: Return ONLY valid JSON in this exact format:
        {{"introduction": "<your introduction text>", "thesis_statement": "<your thesis>"}}"""
        
        intro_response = api_call_with_retry(
            client, job_id, intro_prompt, 
            global_context, # Use updated universal context with image analysis & prompt
            max_tokens=3000
        )
        
        try:
            intro_data = json.loads(intro_response)
        except:
            intro_data = {"introduction": intro_response, "thesis_statement": ""}
        
        update_job_status(job_id, JobStatus.WRITING, 50, f"Generating body sections...")
        
        # Step 3: Generate each body section with adequate word count
        body_sections = []
        for i, section_title in enumerate(sections):
            progress = 50 + int((i / len(sections)) * 30)
            update_job_status(job_id, JobStatus.WRITING, progress, f"Writing section: {section_title}...")
            
            section_prompt = f"""You are an expert academic writer. Write a detailed body section for an academic essay.
            
            Essay Topic: {requirements.get('topic', 'the given topic')}
            Thesis Statement: {intro_data.get('thesis_statement', 'as stated in the introduction')}
            Section Title: {section_title}
            
            CRITICAL LENGTH REQUIREMENT: Write approximately {words_per_section} words. This is essential - the essay must meet word count requirements.
            
            Requirements:
            - Write comprehensive, in-depth analysis with specific examples and evidence
            - Use sophisticated academic language appropriate for {requirements.get('academic_level', 'undergraduate')} level
            - Include clear topic sentences and smooth transitions
            - Develop arguments thoroughly - explain the significance of each point
            - Be thorough and expansive, not brief or superficial
            - Include relevant research, data, or scholarly perspectives where appropriate
            - CRITICAL: Review the "USER ADDITIONAL INSTRUCTIONS" in the provided context and ensure this section aligns with any specific focus areas (e.g. specific country, theory, or case study).
            
            IMPORTANT: Return ONLY valid JSON in this exact format:
            {{"title": "{section_title}", "content": "<your detailed section text>"}}"""
            
            section_response = api_call_with_retry(
                client, job_id, section_prompt,
                global_context, # Updated context
                max_tokens=4000
            )
            
            try:
                section_data = json.loads(section_response)
                body_sections.append({
                    "title": section_data.get("title", section_title),
                    "content": section_data.get("content", ""),
                    "word_count": len(section_data.get("content", "").split())
                })
            except:
                body_sections.append({
                    "title": section_title,
                    "content": section_response,
                    "word_count": len(section_response.split())
                })
        
        update_job_status(job_id, JobStatus.WRITING, 85, "Writing conclusion...")
        
        # Step 4: Generate conclusion
        conclusion_prompt = f"""You are an expert academic writer. Write a strong conclusion for an academic essay.
        
        Topic: {requirements.get('topic', 'the given topic')}
        Thesis Statement: {intro_data.get('thesis_statement', 'as stated')}
        Body Sections Covered: {', '.join(sections)}
        
        Target Length: {conclusion_words} words (CRITICAL: write at least {conclusion_words} words)
        
        Requirements:
        - Synthesize and summarize the key arguments made throughout the essay
        - Restate the thesis in light of the evidence presented
        - Provide meaningful closing thoughts, implications, and future considerations
        - End with a memorable final statement
        - Do NOT introduce new arguments or evidence
        
        IMPORTANT: Return ONLY valid JSON in this exact format:
        {{"conclusion": "<your conclusion text>"}}"""
        
        conclusion_response = api_call_with_retry(
            client, job_id, conclusion_prompt,
            global_context, # Updated context with images/additional prompt
            max_tokens=2000
        )
        
        try:
            conclusion_data = json.loads(conclusion_response)
        except:
            conclusion_data = {"conclusion": conclusion_response}

        # Step 5: Generate References
        update_job_status(job_id, JobStatus.WRITING, 95, "Compiling references...")
        
        references_prompt = f"""You are an expert academic librarian. Compile a list of scholarly references for this essay.
        
        Topic: {requirements.get('topic', 'the given topic')}
        Citation Style: {requirements.get('citation_style', 'APA')}
        
        Requirements:
        - Provide 5-8 relevant, high-quality scholarly sources (journals, books, reputable reports)
        - Format exactly according to {requirements.get('citation_style', 'APA')} style
        - Ensure sources are real and directly relevant to the topic
        
        IMPORTANT: Return ONLY valid JSON in this exact format:
        {{"references": ["<reference 1>", "<reference 2>", ...]}}"""
        
        references_response = api_call_with_retry(
            client, job_id, references_prompt,
            global_context, # Updated context with images/additional prompt
            max_tokens=1500
        )
        
        try:
            references_data = json.loads(references_response)
            references_list = references_data.get("references", [])
        except:
            references_list = []
        
        # Calculate total word count
        total_words = (
            len(intro_data.get("introduction", "").split()) +
            sum(s.get("word_count", 0) for s in body_sections) +
            len(conclusion_data.get("conclusion", "").split())
        )
        
        # Assemble final essay
        essay_data = {
            "title": requirements.get("topic", "Academic Essay"),
            "thesis_statement": intro_data.get("thesis_statement", ""),
            "introduction": intro_data.get("introduction", ""),
            "body_sections": body_sections,
            "conclusion": conclusion_data.get("conclusion", ""),
            "references": references_list,
            "total_word_count": total_words,
            "academic_level": requirements.get("academic_level", "undergraduate")
        }
        
        # Validate with Pydantic
        essay = EssayOutput(**essay_data)
        
        # Store draft
        redis_client.set(f"job:{job_id}:draft", essay.model_dump_json(), ex=86400)
        
        update_job_status(job_id, JobStatus.WRITING, 60, f"Essay generated ({total_words} words)")
        
        # Chain to humanization
        humanize_essay.delay(job_id)
        
        return {"status": "success", "word_count": essay.total_word_count}
        
    except Exception as e:
        update_job_status(job_id, JobStatus.FAILED, 0, error=str(e))
        raise


@celery_app.task(bind=True, name="app.tasks.humanize_essay")
def humanize_essay(self, job_id: str):
    """
    Task to humanize the essay using Burstiness and Perplexity techniques.
    Pass 2: Rewrite the draft to sound more human-written.
    """
    try:
        update_job_status(job_id, JobStatus.HUMANIZING, 70, "Humanizing essay content...")
        
        # Get draft
        draft_data = redis_client.get(f"job:{job_id}:draft")
        if not draft_data:
            raise ValueError("No draft found")
        
        draft = json.loads(draft_data.decode("utf-8") if isinstance(draft_data, bytes) else draft_data)
        
        # Get humanization settings
        job_data = json.loads(redis_client.get(f"job:{job_id}"))
        settings = HumanizationSettings(**job_data.get("humanization_settings", {}))
        additional_prompt = job_data.get("additional_prompt", "")
        
        # Get extracted content for context (including image analysis)
        extracted_content = redis_client.get(f"job:{job_id}:content")
        context_str = extracted_content.decode("utf-8") if extracted_content else ""
        
        from openai import OpenAI
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
        # Separate references to preserve them (GPT might drop them)
        original_references = draft.get("references", [])
        draft_content = {k: v for k, v in draft.items() if k != "references"}
        
        humanization_prompt = f"""Rewrite the following academic essay to sound more naturally 
        human-written while maintaining academic quality. Apply these techniques:
        
        1. BURSTINESS: Vary sentence length significantly. Mix short punchy sentences 
           with longer, more complex ones. Humans don't write uniformly.
        
        2. PERPLEXITY: Use unexpected but appropriate word choices occasionally. 
           Avoid overly predictable phrasing.
        
        3. Natural imperfections: Add minor stylistic variations that feel authentic.
        
        Humanization Intensity: {settings.intensity} (0=minimal, 1=aggressive)
        Preserve Citations: {settings.preserve_citations}
        
        USER ADDITIONAL INSTRUCTIONS (CRITICAL - MUST FOLLOW FOR STYLE/TONE):
        {additional_prompt if additional_prompt else "No additional style instructions."}
        
        Maintain the essay's academic integrity, proper citations, and factual accuracy.
        
        Essay to humanize:
        {json.dumps(draft_content, indent=2)}
        
        Return the humanized essay in the same JSON structure (excluding references)."""
        
        # Retry logic for rate limiting
        max_retries = 5
        for attempt in range(max_retries):
            try:
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "You are an expert editor who humanizes AI-generated text. You MUST strictly follow the user's additional style instructions if provided."},
                        {"role": "user", "content": humanization_prompt}
                    ],
                    response_format={"type": "json_object"},
                    temperature=0.8 + (settings.intensity * 0.2),  # Higher temp for more variation
                    max_tokens=4000
                )
                break
            except Exception as e:
                if "rate_limit" in str(e).lower() or "429" in str(e):
                    wait_time = (2 ** attempt) + random.uniform(0, 1)
                    update_job_status(job_id, JobStatus.HUMANIZING, 75, f"Rate limited, retrying in {int(wait_time)}s...")
                    time.sleep(wait_time)
                    if attempt == max_retries - 1:
                        raise
                else:
                    raise
        
        humanized_json = response.choices[0].message.content
        humanized_data = json.loads(humanized_json)
        
        # Re-attach original references
        humanized_data["references"] = original_references
        
        humanized_essay = EssayOutput(**humanized_data)
        
        # Store humanized version
        redis_client.set(f"job:{job_id}:humanized", humanized_essay.model_dump_json(), ex=86400)
        
        # STOP HERE: Do NOT chain to PDF generation.
        # Set status to WAITING_FOR_REVIEW so user can read/edit.
        update_job_status(job_id, JobStatus.WAITING_FOR_REVIEW, 85, "Ready for review")
        
        return {"status": "success"}
        
    except Exception as e:
        update_job_status(job_id, JobStatus.FAILED, 0, error=str(e))
        raise



@celery_app.task(bind=True, name="app.tasks.refine_essay")
def refine_essay(self, job_id: str, instructions: str):
    """
    Task to refine an existing essay draft based on user feedback.
    """
    try:
        update_job_status(job_id, JobStatus.REFINING, 85, "Refining essay...")
        
        # Get current humanized draft
        essay_data_bytes = redis_client.get(f"job:{job_id}:humanized")
        if not essay_data_bytes:
            raise ValueError("No essay found to refine")
            
        essay_json = essay_data_bytes.decode("utf-8") if isinstance(essay_data_bytes, bytes) else essay_data_bytes
        
        # Calculate current word count for context
        import json
        essay_dict = json.loads(essay_json)
        current_word_count = len(essay_dict.get('introduction', '').split()) + \
                             len(essay_dict.get('conclusion', '').split()) + \
                             sum(len(s.get('content', '').split()) for s in essay_dict.get('body_sections', []))
        
        # Initialize OpenAI client
        from openai import OpenAI
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
        refinement_prompt = f"""You are an expert academic editor.
        
        Refinement Instructions from User:
        "{instructions}"

        Current Essay Stats:
        - REAL Word Count: {current_word_count} words (excluding references).
        
        Current Essay JSON:
        {essay_json}
        
        Task:
        1. If the user asks a question (e.g. "what is the word count?"), answer it in the 'ai_feedback' field and keep the essay content unchanged. USE THE 'REAL WORD COUNT' PROVIDED ABOVE.
        2. If the user asks for edits (e.g., "Expand to 2500 words"), apply them to the essay content content (Introduction, Body Sections, Conclusion). You MUST significantly increase content if asked to expand.
        3. Summarize what you did or answer the question in the 'ai_feedback' field.
        4. Keep the JSON structure EXACTLY the same as the input schema, adding the 'ai_feedback' field at the root level.
        5. Do NOT remove references unless explicitly asked.
        6. Maintain academic tone unless asked to change it.

        Example Output Structure:
        {{
            "title": "...",
            "introduction": "...",
            "body_sections": [...],
            "conclusion": "...",
            "references": [...],
            "ai_feedback": "I have increased the word count by expanding the introduction..."
        }}
        
        Return ONLY the valid updated JSON with the 'ai_feedback' field populated."""
        
        response = api_call_with_retry(
            client, job_id, 
            "You are an intelligent editor. Output valid JSON only.", 
            refinement_prompt, 
            max_tokens=4000
        )
        
        # Validate and store
        updated_data = json.loads(response)
        updated_essay = EssayOutput(**updated_data)
        
        redis_client.set(f"job:{job_id}:humanized", updated_essay.model_dump_json(), ex=86400)
        
        update_job_status(job_id, JobStatus.WAITING_FOR_REVIEW, 85, "Refinement complete")
        return {"status": "success"}

    except Exception as e:
        update_job_status(job_id, JobStatus.FAILED, 0, f"Refinement failed: {str(e)}")
        raise


@celery_app.task(bind=True, name="app.tasks.structure_essay")
def structure_essay(self, job_id: str, raw_text: str, refinement_instructions: Optional[str] = None):
    """
    Takes raw text (from paste or file import) and structures it into
    the EssayOutput JSON format so it can be edited/refined.
    
    If refinement_instructions are provided, it immediately triggers
    a refinement task after structuring.
    """
    try:
        update_job_status(job_id, JobStatus.PLANNING, 10, "Analyzing essay structure...")

        # Initialize OpenAI client
        # (This is safe because we are in a worker process)
        from openai import OpenAI
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

        if not raw_text or len(raw_text.strip()) < 50:
            raise ValueError("Input text is too short to process")

        structuring_prompt = f"""
        You are an Essay Parser.
        
        TASK:
        Take the following "Raw Essay Text" and structure it EXACTLY into the JSON schema provided below.
        
        RULES:
        1. Identify the 'title' (if explicit, or create a simple one based on content).
        2. Identify the 'thesis_statement' (usually the last sentence of the intro, or infer the main argument).
        3. Identify the 'introduction' (first paragraphs).
        4. Split the main body into logical 'body_sections' based on topic changes. Give each section a short 'title' summarizing its topic.
        5. Identify the 'conclusion' (last paragraphs).
        6. Extract the 'references' list. IMPORTANT: Capture the FULL reference entry (Author, Date, Title, Source) exactly as it appears or formatted in APA style. Do NOT just extract in-text citations like (Author, 2023). If no references section exists, return an empty list.
        7. PRESERVE THE ORIGINAL TEXT CONTENT EXACTLY for the body.
        
        Raw Essay Text:
        {raw_text[:15000]}  # Truncate to prevent token overflow if absurdly large, though 15k chars is plenty (~3000 words)

        Output Schema:
        {{
            "title": "string",
            "thesis_statement": "string (one sentence summary of main argument)",
            "introduction": "string (multiline content)",
            "body_sections": [
                {{ "title": "Section Topic", "content": "string (multiline content)" }}
            ],
            "conclusion": "string (multiline content)",
            "references": ["string (Full APA-style citation entry)"]
        }}
        """

        response = api_call_with_retry(
            client=client,
            job_id=job_id,
            system_prompt="You are a strict JSON formatter. Output ONLY valid JSON matching the schema.",
            user_content=structuring_prompt
        )

        # Validate
        import json
        structured_data = json.loads(response)
        
        # Add basic metadata filler if missing
        if "references" not in structured_data:
            structured_data["references"] = []
            
        # Ensure it fits the Pydantic model
        essay_output = EssayOutput(**structured_data)
        
        # Save to Redis as if it were a "humanized" essay ready for review
        redis_client.set(f"job:{job_id}:humanized", essay_output.model_dump_json(), ex=86400)
        
        # Check if user wanted immediate refinement
        if refinement_instructions and len(refinement_instructions.strip()) > 5:
            update_job_status(job_id, JobStatus.REFINING, 20, "Structure complete. Applying initial refinement...")
            # Chain the refinement task
            refine_essay.delay(job_id, refinement_instructions)
        else:
            # Auto-advance to Review stage directly
            update_job_status(job_id, JobStatus.WAITING_FOR_REVIEW, 100, "Import complete")
        
        return {"status": "success", "job_id": job_id}

    except Exception as e:
        update_job_status(job_id, JobStatus.FAILED, 0, f"Import failed: {str(e)}")
        raise


@celery_app.task(bind=True, name="app.tasks.generate_pdf")
def generate_pdf(self, job_id: str):
    """
    Task to generate formatted PDF with academic headers.
    Pass 3: Convert final text to formal PDF.
    """
    try:
        update_job_status(job_id, JobStatus.FORMATTING, 90, "Generating PDF document...")
        
        # Get humanized essay
        essay_data = redis_client.get(f"job:{job_id}:humanized")
        if not essay_data:
            raise ValueError("No humanized essay found")
        
        essay = EssayOutput(**json.loads(essay_data.decode("utf-8") if isinstance(essay_data, bytes) else essay_data))
        
        # Get job metadata
        job_data = json.loads(redis_client.get(f"job:{job_id}"))
        
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        import io
        
        # Create PDF in memory
        pdf_buffer = io.BytesIO()
        
        doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=letter,
            rightMargin=1*inch,
            leftMargin=1*inch,
            topMargin=1*inch,
            bottomMargin=1*inch
        )
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        header_style = ParagraphStyle(
            'Header',
            parent=styles['Normal'],
            fontSize=12,
            alignment=TA_CENTER,
            spaceAfter=6
        )
        
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=12,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            firstLineIndent=0.5*inch
        )
        

        section_style = ParagraphStyle(
            'Section',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=18,
            spaceAfter=12
        )

        reference_style = ParagraphStyle(
            'Reference',
            parent=styles['Normal'],
            fontSize=12,
            leftIndent=0.5*inch,
            firstLineIndent=-0.5*inch,
            spaceAfter=6,
            alignment=TA_JUSTIFY
        )
        
        # Build document content
        story = []
        
        # Academic header
        student_name = job_data.get("student_name") or "Student Name"
        course_name = job_data.get("course_name") or "Course Name"
        current_date = datetime.now().strftime("%B %d, %Y")
        
        story.append(Paragraph(student_name, header_style))
        story.append(Paragraph(course_name, header_style))
        story.append(Paragraph(current_date, header_style))
        story.append(Spacer(1, 0.5*inch))
        
        # Title
        story.append(Paragraph(essay.title, title_style))
        story.append(Spacer(1, 0.25*inch))
        
        import html

        def add_paragraphs(text: str, style):
            """Helper to split text into proper paragraphs."""
            if not text:
                return
            
            # Normalize line endings
            text = text.replace('\r\n', '\n').replace('\r', '\n')
            
            # Split by ANY newline to prevent "Mega Paragraphs" that get truncated
            # ReportLab handles multiple Paragraph objects much better across pages
            # than one giant Paragraph with <br/> tags.
            paragraphs = text.split('\n')
            
            for p in paragraphs:
                if not p.strip():
                    continue
                # Clean and add as separate paragraph
                cleaned = html.escape(p.strip())
                if cleaned:
                    story.append(Paragraph(cleaned, style))

        # Introduction
        add_paragraphs(essay.introduction, body_style)
        
        # Body sections
        for section in essay.body_sections:
            story.append(Paragraph(html.escape(section.title), section_style))
            add_paragraphs(section.content, body_style)
        
        # Conclusion
        story.append(Paragraph("Conclusion", section_style))
        add_paragraphs(essay.conclusion, body_style)
        
        # References
        if essay.references:
            story.append(Spacer(1, 0.5*inch))
             # ... continue with references logic (will need to verify references code too)
            story.append(Paragraph("References", section_style))
            for ref in essay.references:
                story.append(Paragraph(html.escape(ref), reference_style))

        
        doc.build(story)
        
        # Store PDF in Redis
        pdf_bytes = pdf_buffer.getvalue()
        redis_client.set(f"job:{job_id}:pdf", pdf_bytes, ex=86400)
        
        # Generate DOCX version as well
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        docx_doc = Document()
        
        # Add header info
        header_para = docx_doc.add_paragraph()
        header_para.add_run(f"{student_name}\n").bold = False
        header_para.add_run(f"{course_name}\n")
        header_para.add_run(f"{current_date}")
        
        # Add title
        title_para = docx_doc.add_heading(essay.title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Introduction
        docx_doc.add_paragraph(essay.introduction)
        
        # Body sections
        for section in essay.body_sections:
            docx_doc.add_heading(section.title, level=2)
            docx_doc.add_paragraph(section.content)
        
        # Conclusion
        docx_doc.add_heading("Conclusion", level=2)
        docx_doc.add_paragraph(essay.conclusion)
        
        # References
        if essay.references:
            docx_doc.add_heading("References", level=2)
            for ref in essay.references:
                docx_doc.add_paragraph(ref, style='List Bullet')
                
        # Save DOCX to buffer and Redis
        docx_buffer = io.BytesIO()
        docx_doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()
        redis_client.set(f"job:{job_id}:docx", docx_bytes, ex=86400)
        
        # Update final status
        download_url = f"/api/download/{job_id}"
        
        # Update redis job record with download URL
        job_data["status"] = JobStatus.COMPLETED.value
        job_data["progress"] = 100
        job_data["message"] = "Essay generation complete!"
        job_data["download_url"] = download_url
        job_data["updated_at"] = datetime.utcnow().isoformat()
        
        redis_client.set(f"job:{job_id}", json.dumps(job_data), ex=86400)
        
        return download_url
        
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        update_job_status(job_id, JobStatus.FAILED, 0, f"Failed to generate PDF: {str(e)}")
        raise
