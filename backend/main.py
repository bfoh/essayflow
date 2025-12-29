from fastapi import FastAPI, UploadFile, File, HTTPException, Request, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from starlette.middleware.base import BaseHTTPMiddleware
from dotenv import load_dotenv
import uuid
import json
import os
import redis
from datetime import datetime
from typing import Optional
import io
import fitz # PyMuPDF
from docx import Document

# Load environment variables from .env file
load_dotenv()

from app.schemas import (
    JobStatus, 
    TaskStatusResponse, 
    FileUploadResponse,
    JobCreateRequest,
    HumanizationSettings,
    EssayRefinementRequest,
    EssayOutput
)
from app.tasks import process_document, refine_essay, generate_pdf, structure_essay

# Initialize FastAPI app
app = FastAPI(
    title="EssayFlow AI API",
    description="AI-powered academic essay generation service",
    version="1.0.0"
)

# Redis client
redis_client = redis.from_url(os.getenv("REDIS_URL", "redis://localhost:6379/0"))

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "https://essayflo.netlify.app",
        "https://essayflow.netlify.app",
        os.getenv("FRONTEND_URL", "http://localhost:3000")
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# File size limit middleware (10MB max)
class FileSizeLimitMiddleware(BaseHTTPMiddleware):
    MAX_SIZE = 10 * 1024 * 1024  # 10MB in bytes
    
    async def dispatch(self, request: Request, call_next):
        if request.method == "POST" and "upload" in request.url.path:
            content_length = request.headers.get("content-length")
            if content_length:
                if int(content_length) > self.MAX_SIZE:
                    return JSONResponse(
                        status_code=413,
                        content={
                            "detail": f"File size exceeds maximum limit of 10MB"
                        }
                    )
        return await call_next(request)


app.add_middleware(FileSizeLimitMiddleware)


# Upload directory
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "/tmp/essayflow/uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)


@app.get("/")
async def root():
    """Health check endpoint."""
    return {"status": "healthy", "service": "EssayFlow AI API"}


@app.get("/api/health")
async def health_check():
    """Detailed health check endpoint."""
    redis_status = "healthy"
    try:
        redis_client.ping()
    except Exception:
        redis_status = "unhealthy"
    
    return {
        "status": "healthy",
        "redis": redis_status,
        "timestamp": datetime.utcnow().isoformat()
    }


@app.post("/api/upload", response_model=FileUploadResponse)
async def upload_file(
    humanization_intensity: float = Form(0.5),
    student_name: Optional[str] = Form(None),
    course_name: Optional[str] = Form(None),
    additional_prompt: Optional[str] = Form(None),
    file: UploadFile = File(...),
    reference_image_0: Optional[UploadFile] = File(None),
    reference_image_1: Optional[UploadFile] = File(None),
    reference_image_2: Optional[UploadFile] = File(None),
    reference_image_3: Optional[UploadFile] = File(None),
    reference_image_4: Optional[UploadFile] = File(None)
):
    
    """
    Upload a document (PDF or DOCX) to start the essay generation process.
    Also accepts up to 5 reference images.
    """
    # Validate file type
    allowed_types = ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
    allowed_extensions = [".pdf", ".docx"]
    
    file_ext = os.path.splitext(file.filename)[1].lower()
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Only PDF and DOCX files are allowed."
        )
    
    # Generate job ID
    job_id = str(uuid.uuid4())
    
    # Save file (optional, mainly for debugging or backup if needed)
    file_path = os.path.join(UPLOAD_DIR, f"{job_id}{file_ext}")
    
    extracted_text = ""
    
    try:
        contents = await file.read()
        
        # Double-check file size
        if len(contents) > 10 * 1024 * 1024:
            raise HTTPException(
                status_code=413,
                detail="File size exceeds maximum limit of 10MB"
            )
        
        # Save to disk (optional but good for debugging)
        with open(file_path, "wb") as f:
            f.write(contents)
        
        file_size = len(contents)
        
        # Extract text immediately in the API
        import io
        
        if file_ext == ".pdf":
            import fitz  # PyMuPDF
            # Open PDF from bytes
            with fitz.open(stream=contents, filetype="pdf") as doc:
                for page in doc:
                    extracted_text += page.get_text()
                    
        elif file_ext == ".docx":
            from docx import Document
            # Open DOCX from bytes
            with io.BytesIO(contents) as docx_stream:
                doc = Document(docx_stream)
                for para in doc.paragraphs:
                    extracted_text += para.text + "\n"
        
        # Handle Reference Images
        ref_images = [
            img for img in [reference_image_0, reference_image_1, reference_image_2, reference_image_3, reference_image_4]
            if img is not None
        ]
        
        image_count = 0
        for idx, img in enumerate(ref_images):
            try:
                img_content = await img.read()
                if len(img_content) > 0:
                    # Store image in Redis for the worker to process
                    # Use a key like job:{id}:ref_image:{idx}
                    redis_client.set(f"job:{job_id}:ref_image:{idx}", img_content, ex=86400)
                    image_count += 1
            except Exception as e:
                print(f"Failed to save reference image {idx}: {e}")

    except HTTPException:
        raise
    except Exception as e:
        print(f"Error processing file: {str(e)}") # Log error
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process file: {str(e)}"
        )
    
    # Create job record in Redis
    job_data = {
        "job_id": job_id,
        "status": JobStatus.PENDING.value,
        "progress": 0,
        "message": "Job created, waiting to start...",
        "filename": file.filename,
        "file_path": file_path, # Kept for reference
        "file_type": file_ext[1:],
        "created_at": datetime.utcnow().isoformat(),
        "updated_at": datetime.utcnow().isoformat(),
        "student_name": student_name,
        "course_name": course_name,
        "additional_prompt": additional_prompt,
        "ref_image_count": image_count, # Pass count to worker
        "humanization_settings": HumanizationSettings(
            intensity=humanization_intensity
        ).model_dump(),
        "download_url": None,
        "error": None
    }
    
    redis_client.set(f"job:{job_id}", json.dumps(job_data), ex=86400)
    
    # Start processing task - PASS TEXT CONTENT, NOT PATH
    process_document.delay(job_id, extracted_text)
    
    
    return FileUploadResponse(
        job_id=job_id,
        message="File uploaded successfully. Processing started.",
        filename=file.filename,
        file_size=file_size
    )


@app.post("/api/import_essay", response_model=FileUploadResponse)
async def import_essay(
    file: Optional[UploadFile] = File(None),
    text_content: Optional[str] = Form(None),
    refinement_instructions: Optional[str] = Form(None)
):
    """
    Import an existing essay (via file or text paste) for refinement.
    Bypasses generation and goes straight to structuring -> review.
    If refinement_instructions are present, we auto-trigger the refinement task.
    """
    print(f"DEBUG: Processing import request. File: {file.filename if file else 'None'}, Content: {len(text_content) if text_content else 0}")

    if not file and not text_content:
        raise HTTPException(status_code=400, detail="Must provide either a file or text content")

    job_id = str(uuid.uuid4())
    filename = "imported_essay"
    file_size = 0
    extracted_text = ""

    try:
        if file:
            filename = file.filename
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext not in [".pdf", ".docx"]:
                raise HTTPException(status_code=400, detail="Only PDF and DOCX files supported")
            
            print(f"DEBUG: Reading file {filename}...")
            contents = await file.read()
            file_size = len(contents)
            print(f"DEBUG: File read. Size: {file_size} bytes")
            
            if file_ext == ".pdf":
                print("DEBUG: Extracting PDF...")
                with fitz.open(stream=contents, filetype="pdf") as doc:
                    for page_num, page in enumerate(doc):
                        extracted_text += page.get_text()
                print(f"DEBUG: PDF Extracted {len(extracted_text)} chars")
            elif file_ext == ".docx":
                print("DEBUG: Extracting DOCX...")
                with io.BytesIO(contents) as docx_stream:
                    doc = Document(docx_stream)
                    for para in doc.paragraphs:
                        extracted_text += para.text + "\n"
                print(f"DEBUG: DOCX Extracted {len(extracted_text)} chars")
        
        else:
            # Text paste
            print("DEBUG: Using pasted text")
            extracted_text = text_content
            file_size = len(extracted_text.encode('utf-8'))
            filename = "pasted_text.txt"

        if len(extracted_text.strip()) < 50:
             raise HTTPException(status_code=400, detail="Extracted text is too short or empty")

    except Exception as e:
        print(f"ERROR in import_essay: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to process import: {str(e)}")

    # Create job record
    job_data = {
        "job_id": job_id,
        "status": JobStatus.PENDING.value,
        "progress": 0,
        "message": "Importing essay...",
        "filename": filename,
        "created_at": datetime.utcnow().isoformat(),
        "updated_at": datetime.utcnow().isoformat(),
        "student_name": "Student", # Default placeholders
        "course_name": "Course",
        "humanization_settings": {"intensity": 0.5}, # Default
        "download_url": None,
        "error": None
    }
    
    print(f"DEBUG: Setting job {job_id} in Redis")
    redis_client.set(f"job:{job_id}", json.dumps(job_data), ex=86400)
    
    # Store the raw content for reference
    redis_client.set(f"job:{job_id}:content", extracted_text, ex=86400)
    
    # Trigger Structuring Task, passing optional instructions
    print(f"DEBUG: Triggering structure_essay task with instructions: {refinement_instructions} type:{type(refinement_instructions)}")
    structure_essay.delay(job_id, extracted_text, refinement_instructions)
    print("DEBUG: Task triggered.")
    
    return FileUploadResponse(
        job_id=job_id,
        message="Essay imported successfully. Structuring...",
        filename=filename,
        file_size=file_size
    )


@app.get("/api/task/{job_id}", response_model=TaskStatusResponse)
async def get_task_status(job_id: str):
    """
    Get the current status of a job.
    
    Use this endpoint to poll for job progress.
    """
    job_data = redis_client.get(f"job:{job_id}")
    
    if not job_data:
        raise HTTPException(
            status_code=404,
            detail=f"Job with ID '{job_id}' not found"
        )
    
    job = json.loads(job_data)
    
    return TaskStatusResponse(
        job_id=job["job_id"],
        status=JobStatus(job["status"]),
        progress=job["progress"],
        message=job.get("message"),
        created_at=datetime.fromisoformat(job["created_at"]),
        updated_at=datetime.fromisoformat(job["updated_at"]),
        download_url=job.get("download_url"),
        error=job.get("error")
    )


@app.get("/api/download/{job_id}")
async def download_essay(job_id: str, format: str = "pdf"):
    """
    Download the generated essay in PDF or DOCX format.
    
    - **job_id**: The job ID from upload
    - **format**: Output format, either 'pdf' or 'docx' (default: pdf)
    """
    job_data = redis_client.get(f"job:{job_id}")
    
    if not job_data:
        raise HTTPException(
            status_code=404,
            detail=f"Job with ID '{job_id}' not found"
        )
    
    job = json.loads(job_data)
    
    if job["status"] != JobStatus.COMPLETED.value:
        raise HTTPException(
            status_code=400,
            detail=f"Essay is not ready yet. Current status: {job['status']}"
        )
    
    # Determine format and Redis key
    format = format.lower()
    if format == "docx":
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        extension = "docx"
        redis_key = f"job:{job_id}:docx"
    else:
        media_type = "application/pdf"
        extension = "pdf"
        redis_key = f"job:{job_id}:pdf"
    
    # Retrieve file content from Redis
    file_bytes = redis_client.get(redis_key)
    
    if not file_bytes:
        # Fallback to filesystem if not in Redis (for legacy jobs or local dev)
        # But clearly log that we are trying fallback
        file_path = os.path.join(UPLOAD_DIR, f"{job_id}_output.{extension}")
        if os.path.exists(file_path):
             return FileResponse(
                path=file_path,
                media_type=media_type,
                filename=f"{job.get('filename', 'essay')}_output.{extension}"
            )
            
        raise HTTPException(
            status_code=404,
            detail=f"{extension.upper()} file not found (generation might have failed)"
        )
    
    # Get original filename without extension
    original_name = os.path.splitext(job.get("filename", "essay"))[0]
    output_filename = f"{original_name}_essay.{extension}"
    
    from fastapi.responses import Response
    return Response(
        content=file_bytes,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={output_filename}"}
    )


@app.get("/api/essay/{job_id}/content")
async def get_essay_content(job_id: str):
    """
    Get the generated essay content in JSON format.
    
    Useful for displaying in the split-screen editor.
    """
    job_data = redis_client.get(f"job:{job_id}")
    
    if not job_data:
        raise HTTPException(
            status_code=404,
            detail=f"Job with ID '{job_id}' not found"
        )
    
    job = json.loads(job_data)
    
    # Try to get humanized version first, then draft
    essay_data = redis_client.get(f"job:{job_id}:humanized")
    if not essay_data:
        essay_data = redis_client.get(f"job:{job_id}:draft")
    
    if not essay_data:
        raise HTTPException(
            status_code=400,
            detail="Essay content not available yet"
        )
    
    essay = json.loads(essay_data.decode("utf-8") if isinstance(essay_data, bytes) else essay_data)
    
    # Also include original extracted content
    original_content = redis_client.get(f"job:{job_id}:content")
    original_text = ""
    if original_content:
        original_text = original_content.decode("utf-8") if isinstance(original_content, bytes) else original_content
    
    return {
        "job_id": job_id,
        "status": job["status"],
        "essay": essay,
        "original_content": original_text
    }


@app.get("/api/job/{job_id}/review", response_model=EssayOutput)
async def get_essay_for_review(job_id: str):
    """
    Get the humanized essay content for review.
    """
    essay_data = redis_client.get(f"job:{job_id}:humanized")
    if not essay_data:
        raise HTTPException(status_code=404, detail="Essay not ready for review")
    
    return json.loads(essay_data)


@app.post("/api/job/{job_id}/refine")
async def refine_essay_endpoint(job_id: str, request: EssayRefinementRequest):
    """
    Submit instructions to refine the essay.
    """
    # Verify job exists
    if not redis_client.exists(f"job:{job_id}"):
        raise HTTPException(status_code=404, detail="Job not found")
        
    # Trigger refinement task
    refine_essay.delay(job_id, request.instructions)
    
    return {"status": "refining", "message": "Refinement task started"}


@app.post("/api/job/{job_id}/finalize")
async def finalize_essay(job_id: str):
    """
    Approve the essay and generate the final PDF/DOCX.
    """
    # Verify job exists
    if not redis_client.exists(f"job:{job_id}"):
        raise HTTPException(status_code=404, detail="Job not found")
        
    # Trigger PDF generation
    generate_pdf.delay(job_id)
    
    return {"status": "finalizing", "message": "PDF generation started"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
